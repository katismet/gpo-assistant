"""Рендеринг ЛПА в DOCX и PDF."""

import logging
import subprocess
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from io import BytesIO

from docxtpl import DocxTemplate
from docx.shared import Inches

from app.services.bitrix_photos import resize_image_for_document

log = logging.getLogger("gpo.lpa_render")


def render_lpa_docx(
    template_path: str,
    context: Dict[str, Any],
    out_dir: str,
    photos: Optional[List[Tuple[str, bytes]]] = None,
    max_photos_in_doc: int = 5,
) -> Path:
    """Рендерит ЛПА в DOCX с вставкой фото."""
    out_dir_path = Path(out_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)
    
    doc = DocxTemplate(template_path)
    doc.render(context)
    
    # Вставляем фото в конец документа (если есть)
    if photos:
        try:
            from PIL import Image
            
            # Получаем документ после рендеринга
            document = doc.get_docx()
            
            # Добавляем заголовок для фото
            document.add_paragraph("Фото смены:").bold = True
            
            # Вставляем фото (только первые max_photos_in_doc)
            photos_to_insert = photos[:max_photos_in_doc]
            for i, (file_name, image_bytes) in enumerate(photos_to_insert, 1):
                try:
                    # Уменьшаем изображение
                    resized_bytes = resize_image_for_document(image_bytes, max_width=800, max_height=600)
                    
                    # Добавляем параграф с подписью
                    para = document.add_paragraph(f"Фото {i}: {file_name}")
                    para.style = 'List Bullet'
                    
                    # Вставляем изображение
                    image_stream = BytesIO(resized_bytes)
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(5))  # Ширина 5 дюймов
                    
                except Exception as e:
                    log.error(f"Error inserting photo {i}: {e}")
                    # Добавляем текст вместо фото
                    document.add_paragraph(f"Фото {i}: {file_name} (ошибка вставки)")
            
            # Если фото больше max_photos_in_doc, добавляем ссылки на остальные
            if len(photos) > max_photos_in_doc:
                document.add_paragraph(f"\nВсего фото: {len(photos)}. Остальные фото доступны в Bitrix24.")
                
        except ImportError:
            log.warning("PIL not installed, skipping photo insertion")
        except Exception as e:
            log.error(f"Error inserting photos: {e}")
    
    # Имя файла
    safe_obj = (context.get("object_name") or "object").replace("/", "_")
    safe_date = (context.get("date") or "").replace("/", ".").replace("\\", ".")
    out_docx = out_dir_path / f"LPA_{safe_obj}_{safe_date}.docx"
    
    doc.save(str(out_docx))
    return out_docx


def try_convert_pdf(docx_path: Path) -> Optional[Path]:
    """Пытается конвертировать DOCX в PDF через LibreOffice."""
    try:
        out_pdf = docx_path.with_suffix(".pdf")
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            timeout=60,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        return out_pdf if out_pdf.exists() else None
    except FileNotFoundError:
        log.warning("LibreOffice not found, trying docx2pdf")
        try:
            from docx2pdf import convert as docx2pdf_convert
            out_pdf = docx_path.with_suffix(".pdf")
            docx2pdf_convert(str(docx_path), str(out_pdf))
            return out_pdf if out_pdf.exists() else None
        except ImportError:
            log.warning("docx2pdf not installed")
        except Exception as e:
            log.warning(f"docx2pdf conversion failed: {e}")
    except subprocess.TimeoutExpired:
        log.warning("LibreOffice conversion timeout")
    except Exception as e:
        log.warning(f"PDF conversion failed: {e}")
    return None





