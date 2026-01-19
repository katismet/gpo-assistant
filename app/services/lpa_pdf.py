# lpa_pdf.py
import os
import subprocess
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from io import BytesIO
import requests
import logging
import zipfile
import re

from docxtpl import DocxTemplate
from docx.shared import Inches, Cm

log = logging.getLogger("gpo.lpa_pdf")

PLACEHOLDER_PATTERN = re.compile(r"\{\{[^{}]+\}\}")


class LPAPlaceholderError(RuntimeError):
    """Ошибка генерации ЛПА из-за оставшихся плейсхолдеров."""

    def __init__(self, placeholders: List[str], file_path: Path):
        self.placeholders = placeholders
        self.file_path = file_path
        preview = ", ".join(placeholders[:10])
        super().__init__(
            f"Rendered DOCX still contains {len(placeholders)} placeholder(s): {preview}"
        )


# ---- 0) извлечение плейсхолдеров из шаблона ----

def debug_extract_placeholders(template_path: str) -> set:
    """Извлекает все плейсхолдеры вида {{...}} / {% ... %} из всех XML файла DOCX."""
    try:
        from pathlib import Path
        import zipfile
        import re
        
        tpl = Path(template_path)
        if not tpl.exists():
            log.warning(f"[LPA] Template not found for placeholder extraction: {tpl}")
            return set()
        
        placeholders = set()
        patterns = [
            r"\{\{([^}]+)\}\}",
            r"\{%\s*for\s+(\w+)\s+in\s+",
            r"\{%\s*if\s+([^%]+)\s+%\}",
        ]
        
        with zipfile.ZipFile(tpl, "r") as z:
            for name in z.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                try:
                    xml_text = z.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue
                
                for pattern in patterns:
                    for match in re.findall(pattern, xml_text):
                        var_name = match.strip()
                        if "|" in var_name:
                            var_name = var_name.split("|", 1)[0].strip()
                        if var_name and not var_name.startswith("%"):
                            placeholders.add(var_name)
        
        log.info(f"[LPA] Extracted {len(placeholders)} placeholders from template {tpl.name}")
        return placeholders
    except Exception as e:
        log.error(f"[LPA] Error extracting placeholders from template: {e}", exc_info=True)
        return set()


def docx_has_placeholders(docx_path: Path, *, max_examples: int = 50) -> tuple[bool, int, List[str]]:
    """Проверяет все word/*.xml в DOCX на наличие {{...}} плейсхолдеров и возвращает список совпадений."""
    if not docx_path.exists():
        return False, 0, []
    
    count = 0
    examples: List[str] = []
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            for name in z.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                try:
                    xml_text = z.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue
                for match in PLACEHOLDER_PATTERN.finditer(xml_text):
                    count += 1
                    if len(examples) < max_examples:
                        examples.append(match.group(0).replace("\n", ""))
    except Exception as err:
        log.warning(f"[LPA] Could not inspect DOCX for placeholders: {err}")
        return False, 0, []
    return (count > 0, count, examples)

# ---- 1) helpers: нормализация и нарезка под шаблон ----

def _fill_rows(prefix: str, rows: List[Dict[str, Any]], slots: int, mapping: Dict[str, str]) -> Dict[str, Any]:
    """
    prefix='task'| 'equip'| 'worker'| 'mat'
    rows=[{...}, ...]; slots=максимум строк в шаблоне
    mapping: соответствие полей входных данных плейсхолдерам шаблона
    Возвращает dict вида {f"{prefix}{i}_{field}": value}
    """
    ctx = {}
    rows = rows[:slots] if rows else []
    for i in range(1, slots + 1):
        src = rows[i - 1] if i - 1 < len(rows) else {}
        for k_src, k_dst in mapping.items():
            ctx[f"{prefix}{i}_{k_dst}"] = src.get(k_src, "")
    return ctx

def _none(v, default=""):
    return default if v is None else v


def _format_decimal(value: Any, precision: int = 1, suffix: str = "") -> str:
    """Форматирование чисел для шаблона (замена точки на запятую, обрезка нулей)."""
    try:
        num = float(value)
    except (TypeError, ValueError):
        num = 0.0
    fmt = f"{num:.{precision}f}"
    if "." in fmt:
        fmt = fmt.rstrip("0").rstrip(".")
    fmt = fmt.replace(".", ",")
    if not fmt:
        fmt = "0"
    return f"{fmt}{suffix}"


def _flatten_for_template(data: dict) -> dict:
    """Преобразует структурированные данные в плоский контекст для шаблона с пронумерованными плейсхолдерами.
    
    Шаблон ожидает:
    - {{object_name}}, {{object_address}}, {{date}}, {{shift_type}}, {{section}}, {{foreman}}
    - {{task1_name}}, {{task1_unit}}, {{task1_plan}}, {{task1_fact}}, {{task1_executor}}, {{task1_reason}} ... до task10_*
    - {{equip1_name}}, {{equip1_hours}}, {{equip1_comment}} ... до equip7_*
    - {{worker1_name}}, {{worker1_hours}}, {{worker1_rate}}, {{worker1_sum}} ... до worker7_*
    - {{mat1_name}}, {{mat1_unit}}, {{mat1_qty}}, {{mat1_price}}, {{mat1_sum}} ... до mat7_*
    - {{plan_total}}, {{fact_total}}, {{efficiency}}, {{downtime_reason}}, {{reasons_text}}, {{photos_attached}}
    """
    # Начинаем с базовых полей из data
    ctx = {}
    
    # Базовые поля (копируем напрямую)
    base_fields = [
        "object_name", "object_address", "date", "shift_type", 
        "section", "foreman", "downtime_reason", "downtime_min",
        "report_status", "reasons_text"
    ]
    for field in base_fields:
        value = data.get(field)
        if value is None:
            # Для числовых полей используем 0, для строковых - пустую строку
            ctx[field] = "" if field not in ["downtime_min"] else 0
        else:
            ctx[field] = value
    
    ctx["plan_total"] = _format_decimal(data.get("plan_total", 0), precision=1)
    ctx["fact_total"] = _format_decimal(data.get("fact_total", 0), precision=1)
    ctx["efficiency"] = _format_decimal(data.get("efficiency", 0), precision=2, suffix=" %")
    
    # Если object_address нет в data, добавляем пустую строку
    if "object_address" not in ctx:
        ctx["object_address"] = ""
    
    tasks = data.get("tasks", [])
    equips = data.get("tech", []) or data.get("equipment", [])
    workers = data.get("timesheet", [])
    mats = data.get("materials", [])
    
    # Таблица работ: до 10 строк
    for i in range(10):
        t = tasks[i] if i < len(tasks) else {}
        idx = i + 1
        ctx[f"task{idx}_name"] = str(t.get("name", "")).strip()
        ctx[f"task{idx}_unit"] = str(t.get("unit", "")).strip()
        ctx[f"task{idx}_plan"] = _format_decimal(t.get("plan", 0), precision=1)
        ctx[f"task{idx}_fact"] = _format_decimal(t.get("fact", 0), precision=1)
        ctx[f"task{idx}_executor"] = str(t.get("executor", "")).strip()
        ctx[f"task{idx}_reason"] = str(t.get("reason", "")).strip()
    
    # Техника: до 7 строк
    for i in range(7):
        e = equips[i] if i < len(equips) else {}
        idx = i + 1
        ctx[f"equip{idx}_name"] = str(e.get("name", "")).strip()
        ctx[f"equip{idx}_hours"] = _format_decimal(e.get("hours", 0), precision=1) if isinstance(e.get("hours"), (int, float)) else str(e.get("hours", "")).strip()
        ctx[f"equip{idx}_comment"] = str(e.get("comment", "") or e.get("note", "")).strip()
    
    # Табель: до 7 строк
    for i in range(7):
        w = workers[i] if i < len(workers) else {}
        idx = i + 1
        ctx[f"worker{idx}_name"] = str(w.get("name", "")).strip()
        hours_val = w.get("hours", 0) if isinstance(w.get("hours"), (int, float)) else 0
        rate_val = w.get("rate", 0) if isinstance(w.get("rate"), (int, float)) else 0
        ctx[f"worker{idx}_hours"] = _format_decimal(hours_val, precision=1)
        ctx[f"worker{idx}_rate"] = _format_decimal(rate_val, precision=1)
        # Сумма вычисляется, если не задана
        worker_sum = w.get("sum", 0)
        if (not worker_sum) and hours_val and rate_val:
            worker_sum = float(hours_val) * float(rate_val)
        ctx[f"worker{idx}_sum"] = _format_decimal(worker_sum, precision=1)
    
    # Материалы: до 7 строк
    for i in range(7):
        m = mats[i] if i < len(mats) else {}
        idx = i + 1
        ctx[f"mat{idx}_name"] = str(m.get("name", "")).strip()
        ctx[f"mat{idx}_unit"] = str(m.get("unit", "")).strip()
        ctx[f"mat{idx}_qty"] = _format_decimal(m.get("qty", 0), precision=1)
        ctx[f"mat{idx}_price"] = _format_decimal(m.get("price", 0), precision=1)
        # Сумма вычисляется, если не задана
        mat_sum = m.get("sum", 0)
        qty_val = float(m.get("qty", 0) or 0)
        price_val = float(m.get("price", 0) or 0)
        if (not mat_sum) and qty_val and price_val:
            mat_sum = qty_val * price_val
        ctx[f"mat{idx}_sum"] = _format_decimal(mat_sum, precision=1)
    
    # Фото
    photos = data.get("photos", [])
    ctx["photos_attached"] = "Да" if photos else "Нет"
    if photos:
        ctx["photos_attached"] = f"Да ({len(photos)})"
    
    return ctx


def _download_image(url: str) -> Optional[bytes]:
    """Скачивает изображение по URL."""
    try:
        r = requests.get(url, timeout=20)
        if r.ok:
            return r.content
    except Exception as e:
        log.warning(f"Error downloading image from {url}: {e}")
    return None


def _pil_thumb(data: bytes, max_w=800, max_h=600) -> BytesIO:
    """Создает миниатюру изображения."""
    try:
        from PIL import Image
        im = Image.open(BytesIO(data))
        im.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
        out = BytesIO()
        im.save(out, format="JPEG", quality=85)
        out.seek(0)
        return out
    except Exception as e:
        log.warning(f"Error creating thumbnail: {e}")
        return BytesIO(data)


def attach_photos(doc, photos: List[Dict[str, Any]], max_photos_in_doc=5):
    """Вставляет фото в документ.
    
    photos: список dict с ключами:
      - "url" - URL изображения из Bitrix24
      - "tg_file_id" - Telegram file_id (fallback)
    """
    try:
        from PIL import Image
    except ImportError:
        log.warning("PIL not installed, skipping photo insertion")
        return
    
    added = 0
    
    # Добавляем заголовок
    doc.add_paragraph("Фото смены:").bold = True
    
    for p in photos:
        if added >= max_photos_in_doc:
            break
        
        data = None
        photo_name = f"Фото {added + 1}"
        
        # Пробуем скачать из Bitrix24 URL
        if "url" in p:
            url = p["url"]
            data = _download_image(url)
            photo_name = f"Фото {added + 1} (Bitrix24)"
        
        # Fallback к Telegram file_id
        elif "tg_file_id" in p:
            try:
                import os
                from dotenv import load_dotenv
                load_dotenv()
                
                bot_token = os.getenv("BOT_TOKEN")
                if not bot_token:
                    log.warning("BOT_TOKEN not set, cannot download Telegram photo")
                    continue
                
                file_id = p["tg_file_id"]
                # Получаем путь к файлу
                file_info_url = f"https://api.telegram.org/bot{bot_token}/getFile"
                file_info_resp = requests.get(file_info_url, params={"file_id": file_id}, timeout=10)
                if not file_info_resp.ok:
                    log.warning(f"Could not get file info for Telegram file_id {file_id}")
                    continue
                
                file_path = file_info_resp.json().get("result", {}).get("file_path")
                if not file_path:
                    log.warning(f"No file_path for Telegram file_id {file_id}")
                    continue
                
                # Скачиваем файл
                download_url = f"https://api.telegram.org/file/bot{bot_token}/{file_path}"
                download_resp = requests.get(download_url, timeout=30)
                if download_resp.ok:
                    data = download_resp.content
                    photo_name = f"Фото {added + 1} (Telegram)"
                else:
                    log.warning(f"Could not download Telegram file {file_id}")
                    continue
            except Exception as e:
                log.warning(f"Error downloading Telegram photo {file_id}: {e}")
                continue
        
        if not data:
            continue
        
        # Проверяем, что data - это байты, а не строка
        if isinstance(data, str):
            log.warning(f"Photo data is string, not bytes. Skipping photo {added + 1}")
            continue
        
        if not isinstance(data, bytes):
            log.warning(f"Photo data is not bytes (type: {type(data)}). Skipping photo {added + 1}")
            continue
        
        try:
            # Создаем миниатюру
            buff = _pil_thumb(data)
            
            # Добавляем параграф
            para = doc.add_paragraph(photo_name)
            para.style = 'List Bullet'
            
            # Вставляем изображение
            run = para.add_run()
            run.add_picture(buff, width=Inches(4.5))
            
            added += 1
            
        except Exception as e:
            log.error(f"Error inserting photo {added + 1}: {e}")
            doc.add_paragraph(f"{photo_name} (ошибка вставки)")
    
    # Если фото больше max_photos_in_doc, добавляем информацию
    if len(photos) > max_photos_in_doc:
        doc.add_paragraph(f"\nВсего фото: {len(photos)}. Остальные фото доступны в Bitrix24.")

# ---- 2) основной рендер DOCX ----

def render_lpa_docx(
    template_path,
    data: Dict[str, Any],
    out_dir,
    filename_prefix: str = "LPA",
    photos: Optional[List[tuple]] = None,
    max_photos_in_doc: int = 5,
) -> Path:
    """
    data — единый словарь с ключами:
      object_name, section, date, foreman,
      tasks: List[{name, unit, plan, fact, executor, reason}],
      tech/equipment: List[{name, hours, comment}],
      timesheet: List[{name, hours, rate, sum}],
      materials: List[{name, unit, qty, price, sum}],
      plan_total, fact_total, efficiency, downtime_reason,
      photos: List[{url, tg_file_id}]
    photos: Optional[List[tuple]] - устаревший формат, используйте data["photos"]
    max_photos_in_doc: максимальное количество фото для вставки в документ
    Возвращает Path к сгенерированному DOCX.
    """
    # Приводим пути к Path
    tpl = Path(template_path)
    log.info(f"[LPA Render] template_path={tpl} exists={tpl.exists()}")
    if not tpl.exists():
        raise FileNotFoundError(f"Template not found: {tpl}")
    
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Проверяем, что все необходимые поля присутствуют
    object_name = data.get("object_name") or "Не указан"
    tasks = data.get("tasks", [])
    tech = data.get("tech", []) or data.get("equipment", [])
    materials = data.get("materials", [])
    timesheet = data.get("timesheet", [])
    plan_total = data.get("plan_total", 0.0)
    fact_total = data.get("fact_total", 0.0)
    efficiency = data.get("efficiency", 0.0)
    downtime_reason = data.get("downtime_reason", "") or ""
    
    # Преобразуем структурированные данные в плоский контекст для шаблона
    # Шаблон ожидает пронумерованные плейсхолдеры: {{task1_name}}, {{task2_name}}, и т.д.
    flattened_ctx = _flatten_for_template(data)
    
    doc = DocxTemplate(str(tpl))
    
    # Извлекаем плейсхолдеры из шаблона
    template_vars = debug_extract_placeholders(str(tpl))
    context_keys = set(flattened_ctx.keys())
    
    # Сопоставляем плейсхолдеры с контекстом
    missing_in_context = template_vars - context_keys
    missing_in_template = context_keys - template_vars
    
    log.info(f"[LPA] Template vars: {len(template_vars)} items")
    log.info(f"[LPA] Context keys: {len(context_keys)} items")
    log.info(f"[LPA Render] Context keys list: {sorted(context_keys)}")
    
    if missing_in_context:
        log.warning(f"[LPA] Vars only in template: {sorted(missing_in_context)}")
        # Добавляем недостающие переменные в контекст с пустыми значениями
        for var in missing_in_context:
            flattened_ctx[var] = ""
        log.info(f"[LPA] Added {len(missing_in_context)} missing vars to context with empty values")
    
    if missing_in_template:
        log.debug(f"[LPA] Vars only in context: {sorted(missing_in_template)}")
    
    # Рендерим шаблон
    doc.render(flattened_ctx)
    log.info(f"[LPA Render] doc.render() completed successfully")
    
    # Вставляем фото в конец документа (если есть)
    # Используем data["photos"] (список dict с url/tg_file_id)
    photos_list = data.get("photos", [])
    if photos_list:
        log.info(f"[LPA Render] Attaching {len(photos_list)} photos to document")
        rendered_doc = getattr(doc, "docx", None) or doc.get_docx()
        log.info(f"[LPA Render] Using rendered doc type: {type(rendered_doc)}")
        attach_photos(rendered_doc, photos_list, max_photos_in_doc)
        log.info(f"[LPA Render] Photos attached successfully")

    # имя файла
    obj = str(data.get("object_name", "Object")).strip().replace("/", "_")
    date_str = str(data.get("date", "")).strip().replace(":", "_").replace("/", "_")
    name = f"{filename_prefix}_{obj}_{date_str or 'report'}.docx"
    out_docx = out_dir / name
    
    log.info(f"[LPA] Saving final DOCX to: {out_docx}")
    doc.save(str(out_docx))
    if out_docx.exists():
        log.info(
            f"[LPA] Final DOCX saved: {out_docx} (size={out_docx.stat().st_size} bytes)"
        )
    else:
        log.warning(f"[LPA] Final DOCX was not saved: {out_docx}")
    
    has_placeholders, placeholder_count, placeholder_examples = docx_has_placeholders(out_docx)
    log.info(
        f"[LPA] Placeholder scan for {out_docx}: found={has_placeholders}, count={placeholder_count}"
    )
    if has_placeholders:
        if placeholder_examples:
            log.error(
                "[LPA Render] Final DOCX still has %s placeholders. Examples: %s",
                placeholder_count,
                ", ".join(placeholder_examples),
            )
        else:
            log.error(
                "[LPA Render] Final DOCX still has %s placeholders (no sample list available)",
                placeholder_count,
            )
        raise LPAPlaceholderError(placeholder_examples, out_docx)

    log.info("[LPA Render] Final DOCX has no placeholders, safe to convert to PDF")
    log.info(f"[LPA Render] Returning file path: {out_docx}")
    return out_docx

# ---- 3) конвертация DOCX→PDF ----

def docx_to_pdf(input_docx, output_pdf=None, use_libreoffice: bool = True) -> Optional[Path]:
    """
    Пытаемся конвертировать DOCX в PDF:
      A) LibreOffice headless (soffice) — приоритет, если use_libreoffice=True
      B) docx2pdf (хорошо на Windows/macOS) — fallback
      
    Args:
        input_docx: Путь к входному DOCX файлу (str или Path)
        output_pdf: Путь к выходному PDF файлу (опционально, str или Path)
        use_libreoffice: Использовать LibreOffice (True) или docx2pdf (False)
    Returns:
        Path к PDF файлу или None при ошибке
    """
    logger = log

    input_docx = Path(input_docx).resolve()
    if output_pdf is None:
        output_pdf = input_docx.with_suffix(".pdf")
    else:
        output_pdf = Path(output_pdf)
    logger.info(f"[LPA] Converting DOCX to PDF: {input_docx}")

    # Попытка A: LibreOffice headless (приоритет)
    if use_libreoffice:
        try:
            out_dir = str(output_pdf.parent.resolve())
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(input_docx)],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,  # Таймаут 60 секунд
            )
            
            # Проверяем результат
            if output_pdf.exists():
                return output_pdf
            # иногда soffice кладёт рядом с DOCX
            alt = input_docx.with_suffix(".pdf")
            if alt.exists():
                alt.replace(output_pdf)
                return output_pdf
                
            # Если файл не создан, пробуем fallback
            if result.returncode != 0:
                logger.warning(f"LibreOffice conversion failed (return code {result.returncode}), trying fallback")
        except FileNotFoundError:
            # LibreOffice не установлен
            logger.warning("LibreOffice not found, trying docx2pdf")
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice conversion timeout, trying docx2pdf")
        except Exception as e:
            logger.warning(f"LibreOffice conversion error: {e}, trying fallback")

    # Попытка B: docx2pdf (fallback)
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(str(input_docx), str(output_pdf))
        if output_pdf.exists():
            return output_pdf
    except ImportError:
        # docx2pdf не установлен
        pass
    except Exception as e:
        logger.warning(f"docx2pdf conversion failed: {e}")

    # Если оба метода не сработали
    return None

# ---- 4) единая функция: собрать DOCX по шаблону и вернуть PDF ----

def render_lpa_pdf(
    template_path: str,
    out_dir: str,
    data: Dict[str, Any],
    filename_prefix: str = "LPA",
) -> str:
    docx_path = render_lpa_docx(template_path, out_dir, data, filename_prefix)
    pdf_path = docx_to_pdf(docx_path)
    return pdf_path
