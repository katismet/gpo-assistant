#!/usr/bin/env python3
"""
Создание тестового DOCX шаблона для ЛПА
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches

def create_lpa_template():
    """Создаем простой DOCX шаблон для ЛПА"""
    
    # Создаем новый документ
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('ЛИСТ ПРОИЗВОДСТВЕННОГО АНАЛИЗА (ЛПА)', 0)
    title.alignment = 1  # Центрирование
    
    # Основная информация
    doc.add_heading('Основная информация', level=1)
    
    # Таблица с основной информацией
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Объект:'
    table.cell(0, 1).text = '{{object_name}}'
    
    table.cell(1, 0).text = 'Участок:'
    table.cell(1, 1).text = '{{section}}'
    
    table.cell(2, 0).text = 'Дата:'
    table.cell(2, 1).text = '{{date}}'
    
    table.cell(3, 0).text = 'Производитель работ:'
    table.cell(3, 1).text = '{{foreman}}'
    
    # Производственные задания
    doc.add_heading('Производственные задания', level=1)
    
    # Таблица заданий (10 строк)
    tasks_table = doc.add_table(rows=11, cols=7)  # +1 для заголовка, 7 колонок
    tasks_table.style = 'Table Grid'
    
    # Заголовки
    headers = ['№', 'Наименование работ', 'Ед.изм.', 'План', 'Факт', 'Исполнитель', 'Причина отклонения']
    for i, header in enumerate(headers):
        tasks_table.cell(0, i).text = header
    
    # Заполняем плейсхолдеры для заданий
    for i in range(1, 11):
        tasks_table.cell(i, 0).text = str(i)
        tasks_table.cell(i, 1).text = f'{{{{task{i}_name}}}}'
        tasks_table.cell(i, 2).text = f'{{{{task{i}_unit}}}}'
        tasks_table.cell(i, 3).text = f'{{{{task{i}_plan}}}}'
        tasks_table.cell(i, 4).text = f'{{{{task{i}_fact}}}}'
        tasks_table.cell(i, 5).text = f'{{{{task{i}_executor}}}}'
        tasks_table.cell(i, 6).text = f'{{{{task{i}_reason}}}}'
    
    # Техника
    doc.add_heading('Техника', level=1)
    
    equip_table = doc.add_table(rows=8, cols=4)  # +1 для заголовка, 4 колонки
    equip_table.style = 'Table Grid'
    
    # Заголовки техники
    equip_headers = ['№', 'Наименование техники', 'Часы работы', 'Примечания']
    for i, header in enumerate(equip_headers):
        equip_table.cell(0, i).text = header
    
    # Заполняем плейсхолдеры для техники
    for i in range(1, 8):
        equip_table.cell(i, 0).text = str(i)
        equip_table.cell(i, 1).text = f'{{{{equip{i}_name}}}}'
        equip_table.cell(i, 2).text = f'{{{{equip{i}_hours}}}}'
        equip_table.cell(i, 3).text = f'{{{{equip{i}_comment}}}}'
    
    # Табель
    doc.add_heading('Табель', level=1)
    
    timesheet_table = doc.add_table(rows=8, cols=5)  # +1 для заголовка, 5 колонок
    timesheet_table.style = 'Table Grid'
    
    # Заголовки табеля
    timesheet_headers = ['№', 'ФИО', 'Часы', 'Тариф', 'Сумма']
    for i, header in enumerate(timesheet_headers):
        timesheet_table.cell(0, i).text = header
    
    # Заполняем плейсхолдеры для табеля
    for i in range(1, 8):
        timesheet_table.cell(i, 0).text = str(i)
        timesheet_table.cell(i, 1).text = f'{{{{worker{i}_name}}}}'
        timesheet_table.cell(i, 2).text = f'{{{{worker{i}_hours}}}}'
        timesheet_table.cell(i, 3).text = f'{{{{worker{i}_rate}}}}'
        timesheet_table.cell(i, 4).text = f'{{{{worker{i}_sum}}}}'
    
    # Материалы
    doc.add_heading('Материалы', level=1)
    
    materials_table = doc.add_table(rows=8, cols=6)  # +1 для заголовка, 6 колонок
    materials_table.style = 'Table Grid'
    
    # Заголовки материалов
    materials_headers = ['№', 'Наименование', 'Ед.изм.', 'Количество', 'Цена', 'Сумма']
    for i, header in enumerate(materials_headers):
        materials_table.cell(0, i).text = header
    
    # Заполняем плейсхолдеры для материалов
    for i in range(1, 8):
        materials_table.cell(i, 0).text = str(i)
        materials_table.cell(i, 1).text = f'{{{{mat{i}_name}}}}'
        materials_table.cell(i, 2).text = f'{{{{mat{i}_unit}}}}'
        materials_table.cell(i, 3).text = f'{{{{mat{i}_qty}}}}'
        materials_table.cell(i, 4).text = f'{{{{mat{i}_price}}}}'
        materials_table.cell(i, 5).text = f'{{{{mat{i}_sum}}}}'
    
    # Итоговые показатели
    doc.add_heading('Итоговые показатели', level=1)
    
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ('Плановая стоимость:', '{{plan_total}}'),
        ('Фактическая стоимость:', '{{fact_total}}'),
        ('Время простоя (мин):', '{{downtime_min}}'),
        ('Причина простоя:', '{{downtime_reason}}'),
        ('Эффективность (%):', '{{efficiency}}'),
        ('Статус отчёта:', '{{report_status}}')
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
    
    # Дополнительная информация
    doc.add_heading('Дополнительная информация', level=1)
    
    doc.add_paragraph('Причины отклонений:')
    doc.add_paragraph('{{reasons_text}}')
    
    doc.add_paragraph('Фотографии приложены: {{photos_attached}}')
    
    # Сохраняем шаблон
    template_path = Path('app/templates/pdf/lpa_template.docx')
    doc.save(template_path)
    
    print(f"✅ Шаблон ЛПА создан: {template_path}")
    return template_path

if __name__ == "__main__":
    create_lpa_template()
